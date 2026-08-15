import io
import os
import re
from typing import BinaryIO, Union

import fitz
from docx import Document


class ResumeParseError(ValueError):
    """Application-level parsing error for resume documents."""


def _normalize_whitespace(text):
    if not text:
        return ""
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    normalized = re.sub(r"[ \t\f\v]+", " ", normalized)
    normalized = re.sub(r" *\n *", "\n", normalized)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def _safe_filename(filename):
    return os.path.basename(filename or "")


def _extract_pdf_text(data):
    try:
        document = fitz.open(stream=data, filetype="pdf")
    except Exception as exc:
        raise ResumeParseError("The PDF document appears to be corrupted or unreadable.") from exc

    try:
        pages = []
        for page in document:
            text = page.get_text("text") or ""
            if text.strip():
                pages.append(text)
        return "\n\n".join(pages)
    finally:
        try:
            document.close()
        except Exception:
            pass


def _extract_docx_text(data):
    try:
        document = Document(io.BytesIO(data))
    except Exception as exc:
        raise ResumeParseError("The DOCX document appears to be corrupted or unreadable.") from exc

    chunks = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text and cell.text.strip())
            if row_text:
                chunks.append(row_text)

    return "\n".join(chunks)


def _guess_page_count(data, extension):
    if extension == ".pdf":
        try:
            document = fitz.open(stream=data, filetype="pdf")
            page_count = document.page_count
            document.close()
            return page_count
        except Exception:
            return 1
    return 1


def parse_resume_document(source, filename=None):
    """Parse PDF or DOCX resume input into a normalized text payload.

    Accepts either a filesystem path, a bytes payload, or a file-like object.
    """
    if source is None:
        raise ResumeParseError("No resume file was provided.")

    if isinstance(source, str):
        source_name = _safe_filename(source)
        filename = filename or source_name
        if not os.path.exists(source):
            raise ResumeParseError("The uploaded resume could not be found.")
        if not os.path.isfile(source):
            raise ResumeParseError("The uploaded resume is not a valid file.")
        with open(source, "rb") as file_handle:
            data = file_handle.read()
        return _parse_resume_bytes(data, filename)

    if isinstance(source, (bytes, bytearray)):
        return _parse_resume_bytes(bytes(source), filename)

    if hasattr(source, "read"):
        try:
            current_position = source.tell()
        except (AttributeError, OSError):
            current_position = None

        try:
            if hasattr(source, "seek"):
                source.seek(0)
            payload = source.read()
        except Exception as exc:
            raise ResumeParseError("The uploaded resume could not be read.") from exc
        finally:
            if current_position is not None and hasattr(source, "seek"):
                try:
                    source.seek(current_position)
                except Exception:
                    pass

        return _parse_resume_bytes(payload, filename or getattr(source, "filename", None))

    raise ResumeParseError("Unsupported resume document type.")


def _parse_resume_bytes(data, filename=None):
    if data is None or len(data) == 0:
        raise ResumeParseError("The uploaded resume is empty.")

    name = (filename or "").lower()
    extension = os.path.splitext(name)[1].lower()

    if not extension:
        if data.startswith(b"%PDF"):
            extension = ".pdf"
        elif data[:4] == b"PK\x03\x04":
            extension = ".docx"
        else:
            raise ResumeParseError("Unsupported resume file type.")

    if extension == ".pdf":
        extracted_text = _extract_pdf_text(data)
    elif extension == ".docx":
        extracted_text = _extract_docx_text(data)
    else:
        raise ResumeParseError("Unsupported resume file type.")

    normalized_text = _normalize_whitespace(extracted_text)
    if not normalized_text:
        raise ResumeParseError("The uploaded resume is empty.")

    return {
        "text": normalized_text,
        "word_count": len(re.findall(r"\b[\w\-/]+\b", normalized_text)),
        "page_count": _guess_page_count(data, extension),
        "file_extension": extension,
        "filename": _safe_filename(filename),
    }
