import pytest
from docx import Document
import fitz

from app.services.resume_parser import ResumeParseError, parse_resume_document


def _create_pdf(path, text):
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), text)
    doc.save(path)
    doc.close()


def test_parse_normal_pdf(tmp_path):
    file_path = tmp_path / "sample.pdf"
    _create_pdf(file_path, "Alice Student\nComputer Science\nCGPA 3.9\nPython JavaScript")

    result = parse_resume_document(str(file_path))

    assert result["file_extension"] == ".pdf"
    assert "Alice Student" in result["text"]
    assert result["word_count"] > 0
    assert result["page_count"] >= 1


def test_parse_normal_docx(tmp_path):
    file_path = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("Alice Student")
    document.add_paragraph("BSc in Computer Science")
    document.add_paragraph("Skills: Python, Java, SQL")
    document.save(file_path)

    result = parse_resume_document(str(file_path))

    assert result["file_extension"] == ".docx"
    assert "BSc in Computer Science" in result["text"]
    assert "Python" in result["text"]


def test_parse_empty_document(tmp_path):
    empty_pdf = tmp_path / "empty.pdf"
    _create_pdf(empty_pdf, "")
    with pytest.raises(ResumeParseError):
        parse_resume_document(str(empty_pdf))

    empty_docx = tmp_path / "empty.docx"
    Document().save(empty_docx)
    with pytest.raises(ResumeParseError):
        parse_resume_document(str(empty_docx))


def test_parse_invalid_document(tmp_path):
    bad_pdf = tmp_path / "bad.pdf"
    bad_pdf.write_bytes(b"this is not a valid pdf")
    with pytest.raises(ResumeParseError):
        parse_resume_document(str(bad_pdf))

    bad_docx = tmp_path / "bad.docx"
    bad_docx.write_bytes(b"this is not a valid docx")
    with pytest.raises(ResumeParseError):
        parse_resume_document(str(bad_docx))
